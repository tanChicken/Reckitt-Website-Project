"""Generate a Word (.docx) version of the URL routing documentation."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PINK = RGBColor(0xE4, 0x00, 0x46)
NAVY = RGBColor(0x16, 0x21, 0x3E)
GREY = RGBColor(0x55, 0x5B, 0x6B)
CODE_BG_FONT = "Consolas"

doc = Document()

# ---- Base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level > 1 else PINK
    return h


def code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = CODE_BG_FONT
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            if val and (val.startswith("/") or val.startswith("`")):
                run.font.name = CODE_BG_FONT
    doc.add_paragraph()
    return t


# ---- Title ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
trun = title.add_run("URL Routing & Product Pages")
trun.bold = True
trun.font.size = Pt(24)
trun.font.color.rgb = PINK
sub = doc.add_paragraph()
srun = sub.add_run("Reckitt Product Finder — technical change notes")
srun.italic = True
srun.font.color.rgb = GREY
srun.font.size = Pt(11)
doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run(
    "This document describes the routing architecture introduced to make the "
    "product finder shareable and SEO-friendly. Previously the entire experience "
    "lived inside a single client component (ProductFinder) with the current "
    "screen tracked by a local useState step counter — nothing was reflected in "
    "the URL. Now:"
)
bullet("Every wizard screen has a real, shareable link driven by query params.")
bullet("Every product has its own page at /products/<id>.")
bullet("Browser Back / Forward / Refresh all work as expected.")

# ---- Section 1 ----
heading("1. The wizard is driven by the URL", 1)
doc.add_paragraph(
    "The whole funnel state now lives in the query string instead of React state. "
    "This is the single source of truth — ProductFinder reads it on every render "
    "and writes to it on every navigation."
)

heading("Query parameter contract", 2)
table(
    ["Param", "Values", "Meaning"],
    [
        ["step", "welcome | need | symptom | questions | result", "Navigational stage"],
        ["body", "head | throat | heart | stomach", "Body area (NeedSelectionStep)"],
        ["symptom", "sore-throat | cough", "Throat sub-question (throat only)"],
        ["age", "adult | teen | child | someone-else", "Audience"],
        ["severity", "mild | moderate | severe | not-sure", "Symptom severity"],
    ],
)

heading("Example URLs", 2)
table(
    ["URL", "Screen shown"],
    [
        ["/", "Welcome"],
        ["/?step=need", "Body-area selection"],
        ["/?step=symptom&body=throat", "Throat symptom sub-question"],
        ["/?step=questions&body=head", "Age + severity questions"],
        ["/?body=head&age=adult&severity=mild", "Recommendation result"],
        ["/?body=throat&symptom=cough", "Recommendation result (cough -> chest)"],
        ["/?body=heart", "Recommendation result (heart)"],
    ],
)

heading("How the current screen is decided — deriveStep()", 2)
doc.add_paragraph(
    "step is written on every in-app navigation so Back/Forward land exactly where "
    "expected. But it is optional: if it is missing (for example an externally "
    "shared link), the stage is derived from the answers present:"
)
code_block(
    [
        "if no `body`                              -> welcome",
        "if body = throat and no `symptom`         -> symptom",
        "if effective need is heart or chest       -> result   (these skip questions)",
        "if `age` and `severity` are both present  -> result",
        "otherwise                                 -> questions",
    ]
)
doc.add_paragraph(
    "\"Effective need\" accounts for the one rewrite in the flow: a throat body with "
    "a cough symptom is treated as chest for the purposes of the recommendation."
)
doc.add_paragraph(
    "This dual approach means a clean, hand-written link such as "
    "/?body=head&age=adult&severity=mild still resolves straight to the result "
    "page, while internal navigation stays precise."
)

heading("Branching preserved", 2)
bullet("throat -> asks a symptom sub-question first.")
bullet("cough -> rewrites the need to chest and skips the questions screen.")
bullet("heart -> skips the questions screen and goes straight to the result.")
bullet("head / stomach -> go through the questions screen.")

heading("push vs. replace", 2)
doc.add_paragraph(
    "ProductFinder uses a small navigate() helper that writes param updates to the URL:"
)
bullet(
    "Step transitions (Welcome -> Need, Need -> Questions, Next, Back, etc.) use "
    "router.push so each creates a history entry and Back works naturally."
)
bullet(
    "Picking an answer within the questions screen (age / severity) uses "
    "router.replace, so re-selecting an option does not flood the browser history."
)

heading("Suspense requirement", 2)
doc.add_paragraph(
    "useSearchParams() must be wrapped in a <Suspense> boundary in the Next.js App "
    "Router. The home page does this:"
)
code_block(
    [
        "// app/page.tsx",
        "export default function HomePage() {",
        "  return (",
        "    <Suspense fallback={null}>",
        "      <ProductFinder />",
        "    </Suspense>",
        "  );",
        "}",
    ]
)

# ---- Section 2 ----
heading("2. Product pages — /products/<id>", 1)
doc.add_paragraph(
    "Each product in data/productFinder.ts now has its own page, reachable at "
    "/products/<product.id> (e.g. /products/strepsils-original)."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Server component: ").bold = True
p.add_run("app/products/[id]/page.tsx")
bullet("generateStaticParams() pre-renders one static page per product at build time.")
bullet("generateMetadata() sets a per-product title and description for SEO.")
bullet("Unknown ids return notFound() (the 404 page).")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Client island: ").bold = True
p.add_run("components/products/ProductDetail.tsx")
bullet("Handles the interactive variant / flavour selectors.")
bullet(
    "Reuses the same display-derivation logic as the recommendation step, so a "
    "product looks identical whether reached through the finder or via a direct link."
)

heading("Where product links appear", 2)
bullet("Products library (app/products/page.tsx): each card's \"View Details\" button links to /products/<id>.")
bullet("Recommendation step: a \"View full details\" button links to the recommended product's own page.")

# ---- Section 3 ----
heading("3. Files changed", 1)
table(
    ["File", "Change"],
    [
        ["components/product-finder/ProductFinder.tsx", "Rewritten to read/write wizard state from the URL query string."],
        ["app/page.tsx", "Wrapped ProductFinder in a <Suspense> boundary."],
        ["app/products/[id]/page.tsx", "NEW — server route for individual product pages."],
        ["components/products/ProductDetail.tsx", "NEW — client component rendering a single product."],
        ["app/products/page.tsx", "\"View Details\" now links to /products/<id> (was #)."],
        ["components/product-finder/steps/RecommendationStep.tsx", "Added a \"View full details\" link to the product page."],
    ],
)
note = doc.add_paragraph()
note.add_run("Note: ").bold = True
note.add_run(
    "SafetyStep (the former step 5) was already unreachable in the previous code "
    "and is no longer rendered. The component file is kept for future use."
).italic = True

# ---- Section 4 ----
heading("4. How to add a new product", 1)
doc.add_paragraph(
    "1. Add a new entry to productItems in data/productFinder.ts with a unique, "
    "URL-safe id (lowercase, hyphenated — e.g. gaviscon-advance)."
)
doc.add_paragraph(
    "2. Drop the product image at public/products/<id>.png (and any variant/flavour "
    "images using the <id>-<flavor>-<variant>.png convention)."
)
doc.add_paragraph(
    "3. The page at /products/<id> is generated automatically — no routing code needed."
)

# ---- Section 5 ----
heading("5. Testing the changes", 1)
code_block(
    [
        "npm run dev      # then visit the URLs in the tables above",
        "npx tsc --noEmit # type-check",
        "npm run build    # confirms all product pages pre-render as static (SSG)",
    ]
)
doc.add_paragraph(
    "A successful build lists /products/[id] as SSG with one entry per product, "
    "and / as Static."
)

out = "docs/URL-Routing-and-Product-Pages.docx"
doc.save(out)
print("Saved", out)
